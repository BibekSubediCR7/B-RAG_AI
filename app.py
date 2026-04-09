"""
B-RAG AI — Chat with your Documents
Developed by Bibek Subedi
© 2026 All Rights Reserved
"""

# ── Standard library stuff we need ──────────────────────────────────────────
import hashlib      # used to create a unique fingerprint of each uploaded file
import base64       # used to embed images directly into HTML
import requests     # used to talk to Supabase (our database)

# ── Streamlit is the whole UI framework ─────────────────────────────────────
import streamlit as st

# ── Third-party libraries ────────────────────────────────────────────────────
import tiktoken                             # counts tokens so we don't overspend
from pypdf import PdfReader                 # reads and extracts text from PDFs
from docx import Document as DocxDocument   # reads and extracts text from DOCX files
from openai import OpenAI                   # talks to GPT and embedding models
import faiss                                # super fast vector search (finds relevant chunks)
import numpy as np                          # math/array operations for embeddings


# ════════════════════════════════════════════════════════════════════════════
#  PAGE CONFIG — must be the very first Streamlit call, no exceptions
# ════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="B-RAG AI",
    page_icon="assets/favicon.png",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ════════════════════════════════════════════════════════════════════════════
#  SUPABASE SETUP — this is our database that tracks visits, users, etc.
# ════════════════════════════════════════════════════════════════════════════
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]

# These headers go with every request we send to Supabase
HEADERS = {
    "apikey": SUPABASE_KEY,
    "Authorization": f"Bearer {SUPABASE_KEY}",
    "Content-Type": "application/json",
    "Prefer": "return=representation",
}

def _rpc(fn: str, params: dict = {}) -> dict:
    """
    Generic helper to call any Supabase RPC.
    Never crashes the app — if Supabase is slow or down, returns {} silently.
    """
    try:
        r = requests.post(
            f"{SUPABASE_URL}/rest/v1/rpc/{fn}",
            headers=HEADERS,
            json=params,
            timeout=8,       # bumped from 5 — gives Supabase a bit more breathing room
        )
        return r.json() if r.ok else {}
    except requests.exceptions.Timeout:
        return {}            # Supabase timed out — app keeps running normally
    except requests.exceptions.ConnectionError:
        return {}            # no internet / Supabase down — same, keep running
    except Exception:
        return {}            # anything else unexpected — never let stats crash the app


def load_all_stats() -> dict:
    """
    Fetch all four counters in a single DB call.
    One request = visits + users + docs + queries at once.
    Much better than making 4 separate calls.
    """
    result = _rpc("get_all_stats")
    if isinstance(result, list) and result:
        return result[0]
    # Return zeros if something went wrong
    return {"visits": 0, "unique_users": 0, "docs_uploaded": 0, "queries_asked": 0}


def increment_stat(col: str) -> int:
    """
    Add 1 to whichever counter column you name.
    Returns the new value, or 0 if Supabase is unreachable — never crashes.
    """
    try:
        result = _rpc("increment_stat", {"col_name": col})
        return result if isinstance(result, int) else 0
    except Exception:
        return 0


# ── Run this block ONCE per browser session, not on every rerun ─────────────
# Streamlit reruns this whole file constantly. session_state persists between reruns.
if "session_init" not in st.session_state:
    st.session_state["session_init"] = True
    try:
        st.session_state["total_visits"] = increment_stat("visits")
        st.session_state["app_stats"]    = load_all_stats()
        st.session_state["app_stats"]["visits"] = st.session_state["total_visits"]
    except Exception:
        # Supabase completely unreachable at startup — show zeros, keep app running
        st.session_state["total_visits"] = 0
        st.session_state["app_stats"]    = {
            "visits": 0, "docs_uploaded": 0, "queries_asked": 0
        }

# Shortcut to the stats for use elsewhere
VISITS = st.session_state.get("total_visits", 0)


# ════════════════════════════════════════════════════════════════════════════
#  ALL THE CSS STYLING — dark navy + amber theme
# ════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
/* Google Fonts — load the typefaces we use */
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:wght@300;400;500&family=JetBrains+Mono:wght@400;500&display=swap');

/* Color and font variables — change these to retheme the whole app */
:root {
    --navy:      #0B1120;
    --navy-mid:  #121C30;
    --navy-card: #1A2540;
    --amber:     #F5A623;
    --amber-dim: #C4841C;
    --ivory:     #F5F0E8;
    --muted:     #8A95A8;
    --rule-line: #2A3550;
    --success:   #2ECC71;
    --danger:    #E74C3C;
    --font-head: 'Syne', sans-serif;
    --font-body: 'DM Sans', sans-serif;
    --font-mono: 'JetBrains Mono', monospace;
}

/* Base background and text */
html, body, [data-testid="stAppViewContainer"] {
    background-color: var(--navy) !important;
    color: var(--ivory) !important;
    font-family: var(--font-body);
}

[data-testid="stSidebar"] {
    background-color: var(--navy-mid) !important;
    border-right: 1px solid var(--rule-line);
}

/* The small amber badge in the top bar */
.visit-badge {
    position: fixed;
    top: 14px;
    right: 18px;
    background: var(--navy-card);
    border: 1px solid var(--amber);
    color: var(--amber);
    font-family: var(--font-mono);
    font-size: 11px;
    font-weight: 500;
    padding: 4px 12px;
    border-radius: 20px;
    z-index: 9999;
    letter-spacing: 0.05em;
}

/* Big title section at the top of the page */
.hero-block {
    padding: 40px 0 10px 0;
    border-bottom: 2px solid var(--rule-line);
    margin-bottom: 30px;
}
.hero-title {
    font-family: var(--font-head);
    font-size: 52px;
    font-weight: 800;
    color: var(--ivory);
    letter-spacing: -1.5px;
    line-height: 1.1;
    margin: 0;
}
.hero-title span { color: var(--amber); }
.hero-sub {
    font-family: var(--font-body);
    font-size: 16px;
    font-weight: 300;
    color: var(--muted);
    margin-top: 10px;
    letter-spacing: 0.01em;
}

/* Rules / info card with amber left border */
.rules-card {
    background: var(--navy-card);
    border-left: 4px solid var(--amber);
    border-radius: 0 8px 8px 0;
    padding: 20px 24px;
    margin: 24px 0;
}
.rules-card h4 {
    font-family: var(--font-head);
    font-size: 13px;
    font-weight: 700;
    letter-spacing: 0.12em;
    text-transform: uppercase;
    color: var(--amber);
    margin: 0 0 14px 0;
}
.rules-card ul { margin: 0; padding-left: 16px; }
.rules-card ul li {
    font-family: var(--font-body);
    font-size: 14px;
    color: #B0BCCC;
    line-height: 2;
}
.rules-card ul li strong { color: var(--ivory); font-weight: 500; }

/* Red alert card — shown when a file fails validation */
.alert-card {
    background: #1E0A0A;
    border: 1px solid var(--danger);
    border-radius: 8px;
    padding: 18px 22px;
    margin: 16px 0;
}
.alert-card h4 {
    font-family: var(--font-head);
    font-size: 13px;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: var(--danger);
    margin: 0 0 6px 0;
}
.alert-card p { font-size: 14px; color: #E8A0A0; margin: 0; }

/* Green card — shown when a document is processed and ready */
.success-card {
    background: #061A10;
    border: 1px solid var(--success);
    border-radius: 8px;
    padding: 14px 20px;
    font-size: 14px;
    color: #8AEAB8;
    margin: 12px 0;
    font-family: var(--font-mono);
    letter-spacing: 0.03em;
}

/* Small info boxes — used for pages/chunks/tokens in sidebar */
.stat-row { display: flex; gap: 12px; margin: 16px 0; flex-wrap: wrap; }
.stat-chip {
    background: var(--navy-card);
    border: 1px solid var(--rule-line);
    border-radius: 6px;
    padding: 10px 16px;
    font-family: var(--font-mono);
    font-size: 12px;
    color: var(--muted);
}
.stat-chip strong {
    display: block;
    font-size: 18px;
    color: var(--amber);
    font-weight: 600;
    margin-bottom: 2px;
}

/* Chat message bubbles */
[data-testid="stChatMessage"] {
    background: var(--navy-card) !important;
    border: 1px solid var(--rule-line) !important;
    border-radius: 10px !important;
    margin-bottom: 12px !important;
    padding: 14px 18px !important;
}

/* Chat input — bigger, easier to tap on mobile */
[data-testid="stChatInput"] {
    padding: 10px 0 !important;
}
[data-testid="stChatInput"] textarea {
    background: var(--navy-card) !important;
    color: var(--ivory) !important;
    border: 1px solid var(--rule-line) !important;
    border-radius: 12px !important;
    font-family: var(--font-body) !important;
    font-size: 15px !important;
    min-height: 52px !important;
    padding: 14px 16px !important;
    line-height: 1.5 !important;
}
[data-testid="stChatInput"] textarea:focus {
    border-color: var(--amber) !important;
    box-shadow: 0 0 0 3px rgba(245,166,35,0.18) !important;
}
/* Send button — bigger tap target */
[data-testid="stChatInput"] button {
    width: 44px !important;
    height: 44px !important;
    border-radius: 10px !important;
    background: var(--amber) !important;
    border: none !important;
}
[data-testid="stChatInput"] button:hover {
    background: var(--amber-dim) !important;
}

/* File upload drop zone */
[data-testid="stFileUploadDropzone"] {
    background: var(--navy-card) !important;
    border: 2px dashed var(--rule-line) !important;
    border-radius: 12px !important;
    transition: border-color 0.2s;
}
[data-testid="stFileUploadDropzone"]:hover { border-color: var(--amber) !important; }

/* Sidebar profile section */
.profile-name {
    font-family: var(--font-head);
    font-size: 17px;
    font-weight: 700;
    color: var(--ivory);
    text-align: center;
    margin-bottom: 4px;
}
.profile-tagline {
    font-size: 12px;
    color: var(--muted);
    text-align: center;
    line-height: 1.5;
    margin-bottom: 16px;
}

/* Social link buttons in sidebar */
.social-grid { display: flex; flex-direction: column; gap: 8px; margin-top: 12px; }
.social-btn {
    display: block;
    background: var(--navy);
    border: 1px solid var(--rule-line);
    border-radius: 6px;
    padding: 8px 14px;
    font-family: var(--font-body);
    font-size: 12px;
    font-weight: 500;
    color: var(--muted) !important;
    text-decoration: none !important;
    text-align: center;
    transition: all 0.2s;
    letter-spacing: 0.04em;
    cursor: pointer;
}
.social-btn:hover {
    border-color: var(--amber);
    color: var(--amber) !important;
    background: rgba(245,166,35,0.06);
}

/* Thin horizontal line used as a section separator in sidebar */
.sidebar-divider { border: none; border-top: 1px solid var(--rule-line); margin: 20px 0; }

/* Small uppercase label above each sidebar section */
.sidebar-section-label {
    font-family: var(--font-head);
    font-size: 10px;
    font-weight: 700;
    letter-spacing: 0.14em;
    text-transform: uppercase;
    color: var(--amber);
    margin-bottom: 12px;
    display: block;
}

/* Footer at the very bottom */
.footer-bar {
    text-align: center;
    padding: 40px 0 20px 0;
    border-top: 1px solid var(--rule-line);
    margin-top: 60px;
    font-family: var(--font-mono);
    font-size: 11px;
    color: var(--muted);
    letter-spacing: 0.06em;
}

/* Minor Streamlit overrides */
.stSpinner > div { color: var(--amber) !important; }
div[data-testid="stMarkdownContainer"] p { color: var(--ivory); }
button[kind="primary"] {
    background: var(--amber) !important;
    color: var(--navy) !important;
    font-family: var(--font-head) !important;
    font-weight: 700 !important;
    border: none !important;
    border-radius: 6px !important;
}
button[kind="secondary"] {
    background: transparent !important;
    color: var(--amber) !important;
    border: 1px solid var(--amber) !important;
    font-family: var(--font-head) !important;
    font-weight: 600 !important;
    border-radius: 6px !important;
}
</style>
<script>
// After user submits — blur the input so mobile keyboard closes
// and scroll down so they can see the response
const observer = new MutationObserver(() => {
    const textarea = document.querySelector('[data-testid="stChatInput"] textarea');
    if (!textarea) return;

    textarea.addEventListener('keydown', function(e) {
        // Enter without Shift = submit
        if (e.key === 'Enter' && !e.shiftKey) {
            setTimeout(() => {
                // Blur closes the keyboard on iOS/Android
                textarea.blur();
                // Scroll to bottom so response is visible
                window.scrollTo({ top: document.body.scrollHeight, behavior: 'smooth' });
            }, 150);
        }
    });

    // Also handle the send button tap on mobile
    const btn = document.querySelector('[data-testid="stChatInput"] button');
    if (btn) {
        btn.addEventListener('click', () => {
            setTimeout(() => {
                textarea.blur();
                window.scrollTo({ top: document.body.scrollHeight, behavior: 'smooth' });
            }, 150);
        });
    }
});

observer.observe(document.body, { childList: true, subtree: true });
</script>
""", unsafe_allow_html=True)




# ════════════════════════════════════════════════════════════════════════════
#  OPENAI CLIENT — single instance, used everywhere
# ════════════════════════════════════════════════════════════════════════════
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])


# ════════════════════════════════════════════════════════════════════════════
#  CONSTANTS — tweak these to control cost and quality
# ════════════════════════════════════════════════════════════════════════════
EMBED_MODEL       = "text-embedding-3-small"  # cheapest OpenAI embedding model
CHAT_MODEL        = "gpt-4o-mini"             # cheapest good chat model
MAX_TOKENS        = 800                        # max words GPT can reply with
MAX_CTX_TOKS      = 4000                       # max tokens we send as context per question
MAX_FILE_MB       = 10                         # reject files bigger than this
MAX_PAGES         = 300                        # reject PDFs longer than this
CHUNK_SIZE        = 400                        # how many words per chunk (was 500 — reduced to save tokens)
CHUNK_OVERLAP     = 40                         # how many words overlap between chunks
EMBED_DIM         = 1536                       # dimension of text-embedding-3-small vectors
TOP_K_CHUNKS      = 3                          # how many chunks to retrieve per question (was 5 — fewer = cheaper)
MAX_HISTORY_TURNS = 6                          # only keep last 6 Q&A pairs in memory — stops history growing forever


# The personality and instructions we give GPT at the start of every chat
SYSTEM_PROMPT = SYSTEM_PROMPT = """You are an expert analyst with full mastery of the provided document. Deliver precise, confident analysis — never describe what the document says, just answer.

Format: Markdown always. Bold key terms. ## Headers for sections. Bullet points for lists. Tables for numbers/comparisons. 3-4 paragraphs max per answer.

Tone: Domain expert. No disclaimers like "based on the text." If something is outside the document scope, say so briefly."""

# ════════════════════════════════════════════════════════════════════════════
#  SESSION STATE — these variables survive across Streamlit reruns
# ════════════════════════════════════════════════════════════════════════════
for key, default in {
    "vector_store":  None,   # the FAISS index (our search engine for the doc)
    "chunks":        [],     # list of text chunks from the document
    "chat_history":  [],     # all messages sent and received so far
    "doc_stats":     {},     # info about the current doc (pages, chunks, tokens)
    "file_name":     None,   # name of the currently loaded file
    "doc_hash":      None,   # MD5 fingerprint of the file (used for embedding cache)
    "uploader_key":  0,      # increment this to force the uploader to reset
}.items():
    if key not in st.session_state:
        st.session_state[key] = default


# ════════════════════════════════════════════════════════════════════════════
#  CACHED RESOURCES — @st.cache_resource means created ONCE and reused forever
#                     @st.cache_data means result is stored and returned if
#                     the same inputs come in again — no re-running the function
# ════════════════════════════════════════════════════════════════════════════

@st.cache_resource
def get_tokenizer():
    """
    Load the token counter once and reuse it forever.
    Without this, tiktoken would reload its vocab file on every single call.
    """
    return tiktoken.encoding_for_model("gpt-4o-mini")


@st.cache_data(show_spinner=False)
def get_base64_image(path: str) -> str:
    """
    Read an image file and convert it to a base64 string.
    Cached so we don't re-read the file from disk on every Streamlit rerun.
    """
    try:
        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except FileNotFoundError:
        return ""


@st.cache_data(show_spinner=False)
def get_embeddings_cached(cache_key: str, texts: tuple[str, ...]) -> np.ndarray:
    """
    THE MOST IMPORTANT CACHING FUNCTION.

    Converts text into embedding vectors using OpenAI's API.
    The cache_key is an MD5 hash — if we've seen these exact texts before,
    Streamlit returns the stored result immediately with ZERO API calls.

    This is what was burning your 673k embedding tokens — the same document
    was being re-embedded every time because there was no cache.

    texts must be a tuple (not a list) because Streamlit can only cache
    hashable types, and lists are not hashable.
    """
    response = client.embeddings.create(model=EMBED_MODEL, input=list(texts))
    return np.array([item.embedding for item in response.data], dtype=np.float32)


def get_embeddings(texts: list[str], cache_key: str | None = None) -> np.ndarray:
    """
    Public wrapper around get_embeddings_cached.
    If you don't provide a cache_key, we generate one by hashing the text itself.
    This means: same question asked twice = zero API calls the second time.
    """
    if cache_key is None:
        # Hash the content of the text to create a stable cache key
        cache_key = hashlib.md5("".join(texts).encode()).hexdigest()
    return get_embeddings_cached(cache_key, tuple(texts))


@st.cache_data(show_spinner=False)
def build_faiss_index_cached(doc_hash: str, chunks: tuple[str, ...]):
    """
    Build the FAISS search index for a document.
    Keyed by the file's MD5 hash — so if someone uploads the SAME file
    in a brand new session, we return the cached index with zero embedding calls.

    This is the single biggest fix for your token bill.
    """
    embeddings = get_embeddings(list(chunks), cache_key=doc_hash)
    faiss.normalize_L2(embeddings)                    # normalize for cosine similarity
    index = faiss.IndexFlatIP(EMBED_DIM)              # Inner Product = cosine after normalization
    index.add(embeddings)
    return index, embeddings


# ════════════════════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ════════════════════════════════════════════════════════════════════════════

def security_alert(message: str):
    """Show a red alert box with a custom message."""
    st.markdown(f"""
    <div class="alert-card">
        <h4>Security Alert</h4>
        <p>{message}</p>
    </div>
    """, unsafe_allow_html=True)


def file_hash(data: bytes) -> str:
    """
    Create an MD5 fingerprint from raw file bytes.
    Two identical files will always produce the same hash.
    This is what lets us cache embeddings across sessions.
    """
    return hashlib.md5(data).hexdigest()


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """
    Split a long text into smaller overlapping pieces (chunks).
    Overlap means consecutive chunks share some words — so a sentence
    that falls at a chunk boundary doesn't get split and lost.
    """
    words = text.split()
    chunks, i = [], 0
    while i < len(words):
        chunk = " ".join(words[i : i + size])
        chunks.append(chunk)
        i += size - overlap
    return chunks


def retrieve(query: str, index, chunks: list[str], k: int = TOP_K_CHUNKS) -> list[str]:
    """
    Find the k chunks most relevant to the user's question.
    We embed the question, then ask FAISS which stored embeddings are closest.
    Query embedding is auto-cached — same question twice = no extra API call.
    """
    q_emb = get_embeddings([query])          # cached by query text hash
    faiss.normalize_L2(q_emb)
    _, indices = index.search(q_emb, k)
    return [chunks[i] for i in indices[0] if i < len(chunks)]


def trim_context_to_budget(chunks: list[str], budget: int = MAX_CTX_TOKS) -> str:
    """
    Join the retrieved chunks into one context string.
    Stop adding chunks if we'd go over our token budget.
    This is the cost guardrail — we never send more than MAX_CTX_TOKS tokens to GPT.
    """
    enc = get_tokenizer()          # uses cached tokenizer, not a fresh load
    context, tokens_used = [], 0
    for chunk in chunks:
        chunk_toks = len(enc.encode(chunk))
        if tokens_used + chunk_toks > budget:
            break
        context.append(chunk)
        tokens_used += chunk_toks
    return "\n\n---\n\n".join(context)


def trim_chat_history(history: list[dict], max_turns: int = MAX_HISTORY_TURNS) -> list[dict]:
    """
    Keep only the last N question-answer pairs in memory.
    Without this, a long chat session would keep growing the context window
    and silently spend more and more tokens on every request.
    Each turn = 1 user message + 1 assistant message = 2 items in the list.
    """
    max_items = max_turns * 2
    return history[-max_items:] if len(history) > max_items else history


def estimate_tokens(text: str) -> int:
    """Count how many tokens a string is — uses cached tokenizer."""
    return len(get_tokenizer().encode(text))


def count_extractable_chars(reader: PdfReader) -> int:
    """
    Count total characters we can extract from a PDF.
    Low character count usually means the PDF is a scanned image, not real text.
    """
    total = 0
    for page in reader.pages:
        total += len(page.extract_text() or "")
    return total


def extract_pdf_text(reader: PdfReader) -> str:
    """Pull all text out of a PDF, one page at a time."""
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n".join(parts)


def extract_docx_text(file_bytes: bytes) -> str:
    """
    Pull all text out of a DOCX file — paragraphs and tables.
    We also grab table content so tabular data isn't silently lost.
    """
    import io
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = []
    # Get all normal paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Also grab table cells (joined with pipe separators)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


# ════════════════════════════════════════════════════════════════════════════
#  SIDEBAR
# ════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    _s = st.session_state.get("app_stats", {})
    st.markdown(f"""
    <div style="display:grid; grid-template-columns:repeat(3,1fr); gap:5px; margin-bottom:10px;">
        <div style="background:var(--navy-card); border:1px solid var(--rule-line); border-radius:6px; padding:5px 4px; text-align:center;">
            <div style="font-family:var(--font-mono); font-size:13px; font-weight:700; color:var(--amber); line-height:1.2;">{_s.get('visits', 0):,}</div>
            <div style="font-family:var(--font-body); font-size:9px; color:var(--muted); letter-spacing:0.04em;">VISITS</div>
        </div>
        <div style="background:var(--navy-card); border:1px solid var(--rule-line); border-radius:6px; padding:5px 4px; text-align:center;">
            <div style="font-family:var(--font-mono); font-size:13px; font-weight:700; color:var(--amber); line-height:1.2;">{_s.get('docs_uploaded', 0):,}</div>
            <div style="font-family:var(--font-body); font-size:9px; color:var(--muted); letter-spacing:0.04em;">DOCS</div>
        </div>
        <div style="background:var(--navy-card); border:1px solid var(--rule-line); border-radius:6px; padding:5px 4px; text-align:center;">
            <div style="font-family:var(--font-mono); font-size:13px; font-weight:700; color:var(--amber); line-height:1.2;">{_s.get('queries_asked', 0):,}</div>
            <div style="font-family:var(--font-body); font-size:9px; color:var(--muted); letter-spacing:0.04em;">QUERIES</div>
        </div>
    </div>
    <hr style="border:none; border-top:1px solid var(--rule-line); margin:8px 0 10px 0;">
    """, unsafe_allow_html=True)
    # ── DEVELOPER PROFILE ────────────────────────────────────────────────
    st.markdown('<span class="sidebar-section-label">Developer Profile</span>', unsafe_allow_html=True)
    st.image("assets/photo.jpg", use_container_width=True)
    st.markdown('<div class="profile-name">Bibek Subedi</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="profile-tagline">Aspiring AI Researcher<br>& Data Scientist</div>',
        unsafe_allow_html=True,
    )
    st.markdown("""
    <div class="social-grid">
        <a class="social-btn" href="https://github.com/BibekSubediCR7" target="_blank">GitHub &rarr;</a>
        <a class="social-btn" href="https://www.linkedin.com/in/bibeksubedicr7/" target="_blank">LinkedIn &rarr;</a>
        <a class="social-btn" href="https://www.facebook.com/profile.php?id=100015784387352" target="_blank">Facebook &rarr;</a>
        <a class="social-btn" href="https://www.instagram.com/bibek_subedi_cr7" target="_blank">Instagram &rarr;</a>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<hr class="sidebar-divider">', unsafe_allow_html=True)

    # ── CURRENT DOCUMENT STATS ────────────────────────────────────────────
    # Only shown after a document is processed
    st.markdown('<span class="sidebar-section-label">App Controls</span>', unsafe_allow_html=True)

    if st.session_state["doc_stats"]:
        stats = st.session_state["doc_stats"]
        st.markdown(f"""
        <div class="stat-row">
            <div class="stat-chip"><strong>{stats.get('pages', '—')}</strong>Pages</div>
            <div class="stat-chip"><strong>{stats.get('chunks', '—')}</strong>Chunks</div>
            <div class="stat-chip"><strong>{stats.get('tokens', '—'):,}</strong>Tokens</div>
        </div>
        """, unsafe_allow_html=True)

    # Clear button — resets the document and chat but keeps the app running
    if st.session_state["vector_store"] is not None:
        if st.button("Clear Document & Reset", type="secondary", use_container_width=True):
            for k in ["vector_store", "chunks", "chat_history", "doc_stats", "file_name", "doc_hash"]:
                st.session_state[k] = [] if k in ["chunks", "chat_history"] else (None if k != "doc_stats" else {})
            st.rerun()

    st.markdown('<hr class="sidebar-divider">', unsafe_allow_html=True)

    # ── MODEL INFO ────────────────────────────────────────────────────────
    st.markdown('<span class="sidebar-section-label">Model Configuration</span>', unsafe_allow_html=True)
    st.markdown(f"""
    <div style="font-family: var(--font-mono); font-size: 11px; color: var(--muted); line-height: 2;">
        Chat &nbsp;&nbsp;&nbsp;&nbsp;: {CHAT_MODEL}<br>
        Embed &nbsp;&nbsp;&nbsp;: {EMBED_MODEL}<br>
        Max Tok : {MAX_TOKENS}<br>
        Ctx Tok &nbsp;: {MAX_CTX_TOKS}<br>
        Top-K &nbsp;&nbsp;&nbsp;: {TOP_K_CHUNKS}
    </div>
    """, unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════════════════
#  HERO BANNER — big title with background image at the top
# ════════════════════════════════════════════════════════════════════════════

# Load the cover image as base64 — cached so it's only read from disk once
_img_b64 = get_base64_image("assets/cover.jpg")
bg_css = f"url('data:image/jpeg;base64,{_img_b64}')" if _img_b64 else "none"

st.markdown(f"""
<div class="hero-block" style="
    background-image: {bg_css};
    background-size: cover;
    background-position: center;
    border-radius: 12px;
    padding: 60px 40px;
    position: relative;
    overflow: hidden;
    margin-bottom: 30px;
">
    <div style="
        position: absolute;
        inset: 0;
        background: linear-gradient(
            to right,
            rgba(11,17,32,0.92) 40%,
            rgba(11,17,32,0.55) 100%
        );
        border-radius: 12px;
    "></div>
    <div style="position: relative; z-index: 1;">
        <h1 class="hero-title">B-<span>RAG</span> AI</h1>
        <p class="hero-sub">An intelligent RAG system that understands your documents.</p>
    </div>
</div>
""", unsafe_allow_html=True)


# ── Document rules shown to the user before they upload ─────────────────────
st.markdown("""
<div class="rules-card">
    <h4>Document Submission Policy</h4>
    <ul>
        <li><strong>File size</strong> must not exceed <strong>10 MB</strong>.</li>
        <li><strong>Page count</strong> must not exceed <strong>300 pages</strong>.</li>
        <li>Documents must be <strong>text-based</strong>. Scanned image PDFs/DOCXs are not supported.</li>
        <li>Supported formats: <strong>PDF</strong> and <strong>DOCX</strong>.</li>
    </ul>
</div>
""", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════════════════
#  FILE UPLOAD — accepts both PDF and DOCX now
# ════════════════════════════════════════════════════════════════════════════

# uploader_key is incremented when the user clicks Clear — this forces
# Streamlit to completely destroy and recreate the uploader widget (clearing it visually)
uploaded_file = st.file_uploader(
    "Upload PDF or DOCX Document",
    type=["pdf", "docx"],          # DOCX support added here
    label_visibility="collapsed",
    key=f"doc_uploader_{st.session_state['uploader_key']}",
)

# Clear button — resets everything including the visual file picker
if st.button("Clear"):
    st.session_state["uploader_key"] += 1
    st.session_state["vector_store"] = None
    st.session_state["file_name"]    = None
    st.session_state["doc_hash"]     = None
    st.session_state["chunks"]       = []
    st.session_state["chat_history"] = []
    st.session_state["doc_stats"]    = {}
    st.rerun()


# ════════════════════════════════════════════════════════════════════════════
#  FILE VALIDATION & PROCESSING
# ════════════════════════════════════════════════════════════════════════════
if uploaded_file is not None:

    # ── Rule 1: File size check ──────────────────────────────────────────
    file_size_mb = uploaded_file.size / (1024 * 1024)
    if file_size_mb > MAX_FILE_MB:
        security_alert(
            f"File size ({file_size_mb:.1f} MB) exceeds the 10 MB limit. "
            f"Please click Clear and upload a smaller file."
        )

    else:
        # Figure out what kind of file this is
        is_docx = uploaded_file.name.lower().endswith(".docx")

        if is_docx:
            # DOCX files skip PDF-specific checks (page count, text density)
            reader         = None
            page_count     = 1    # placeholder — recalculated during extraction
            chars_per_page = 999  # always passes the text-density check
        else:
            # PDF — run all three validation rules
            reader     = PdfReader(uploaded_file)
            page_count = len(reader.pages)

        # ── Rule 2: Page count (PDF only) ───────────────────────────────
        if not is_docx and page_count > MAX_PAGES:
            security_alert(
                f"Document has {page_count} pages, exceeding the {MAX_PAGES}-page limit. "
                f"Please click Clear and upload a shorter document."
            )

        else:
            # ── Rule 3: Text density (PDF only) ─────────────────────────
            # A very low chars-per-page ratio usually means it's a scanned image PDF
            if not is_docx:
                char_count     = count_extractable_chars(reader)
                chars_per_page = char_count / max(page_count, 1)

            if not is_docx and chars_per_page < 50:
                security_alert(
                    "This PDF appears to be a scanned image without selectable text. "
                    "Please click Clear and upload a text-based PDF."
                )

            # ── All rules passed — now process the document ──────────────
            else:
                # Read the raw bytes once — we need them for hashing and DOCX parsing
                file_bytes = uploaded_file.getvalue()
                doc_hash   = file_hash(file_bytes)

                # Check if this is a different file than what's already loaded
                is_new_file = st.session_state.get("file_name") != uploaded_file.name

                if is_new_file:
                    # Reset everything if a new file is uploaded
                    st.session_state["vector_store"] = None
                    st.session_state["chunks"]       = []
                    st.session_state["chat_history"] = []
                    st.session_state["doc_stats"]    = {}
                    st.session_state["file_name"]    = uploaded_file.name
                    st.session_state["doc_hash"]     = doc_hash

                if st.session_state["vector_store"] is None:
                    with st.spinner("Analyzing document..."):

                        # ── Extract text based on file type ─────────────
                        if is_docx:
                            full_text  = extract_docx_text(file_bytes)
                            # For DOCX, count double-newlines as a proxy for "sections"
                            page_count = max(1, full_text.count("\n\n"))
                        else:
                            full_text = extract_pdf_text(reader)

                        # Split into chunks
                        chunks = chunk_text(full_text)

                        # Build FAISS index — CACHED by file hash
                        # If someone uploads the same file again (even in a new session),
                        # this returns immediately with zero OpenAI API calls
                        index, _ = build_faiss_index_cached(doc_hash, tuple(chunks))

                        # Store everything in session state
                        st.session_state["chunks"]       = chunks
                        st.session_state["vector_store"] = index
                        st.session_state["doc_stats"]    = {
                            "pages":  page_count,
                            "chunks": len(chunks),
                            "tokens": estimate_tokens(full_text),
                        }

                        # Track this document upload in Supabase
                        if is_new_file:
                            new_count = increment_stat("docs_uploaded")
                            if "app_stats" in st.session_state:
                                st.session_state["app_stats"]["docs_uploaded"] = new_count

                    st.rerun()

                # ── Show green success card when document is ready ───────
                stats = st.session_state["doc_stats"]
                st.markdown(f"""
                <div class="success-card">
                    ✓ READY &nbsp;|&nbsp; {uploaded_file.name} &nbsp;|&nbsp;
                    {stats.get('pages')} sections /
                    {stats.get('chunks')} chunks /
                    ~{stats.get('tokens'):,} tokens
                </div>
                """, unsafe_allow_html=True)

                st.markdown("---")


# ════════════════════════════════════════════════════════════════════════════
#  CHAT INTERFACE — only shown after a document is loaded
# ════════════════════════════════════════════════════════════════════════════
if st.session_state["vector_store"] is not None:

    # ── Show all previous messages in this session ───────────────────────
    for msg in st.session_state["chat_history"]:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # ── Wait for the user to type a question ────────────────────────────
    user_input = st.chat_input("Ask anything about your document...")

    if user_input:
        # Show the user's message immediately
        st.session_state["chat_history"].append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.markdown(user_input)

        # ── Find the most relevant chunks for this question ──────────────
        # Query embedding is cached — repeated questions cost nothing
        relevant_chunks = retrieve(
            user_input,
            st.session_state["vector_store"],
            st.session_state["chunks"],
        )

        # Trim the chunks to our token budget before sending to GPT
        context = trim_context_to_budget(relevant_chunks)

        # Trim chat history so old messages don't silently inflate token count
        st.session_state["chat_history"] = trim_chat_history(st.session_state["chat_history"])

        # ── Build the message list we send to GPT ───────────────────────
        # We only send: system prompt + the relevant context + the question.
        # We do NOT send the full chat history to GPT — just the document context.
        # This keeps every request flat and predictable in token cost.
        messages_payload = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": (
                    f"**Context from document:**\n\n{context}\n\n"
                    f"---\n\n**Question:** {user_input}"
                ),
            },
        ]

        # Track this query in Supabase — updates the Queries counter
        new_q = increment_stat("queries_asked")
        if "app_stats" in st.session_state:
            st.session_state["app_stats"]["queries_asked"] = new_q

        # ── Stream the response word by word ────────────────────────────
        with st.chat_message("assistant"):
            response_placeholder = st.empty()
            full_response = ""

            stream = client.chat.completions.create(
                model       = CHAT_MODEL,
                messages    = messages_payload,
                max_tokens  = MAX_TOKENS,
                temperature = 0.3,      # low temperature = more focused, less random
                stream      = True,
            )

            # Each chunk arrives one at a time — we concatenate and re-render
            for chunk in stream:
                delta = chunk.choices[0].delta.content or ""
                full_response += delta
                response_placeholder.markdown(full_response + "▌")  # blinking cursor effect

            # Final render without the cursor
            response_placeholder.markdown(full_response)

        # Save the assistant's reply to history
        st.session_state["chat_history"].append(
            {"role": "assistant", "content": full_response}
        )

else:
    # ── Placeholder shown when no document has been uploaded yet ─────────
    st.markdown("""
    <div style="
        text-align: center;
        padding: 60px 20px;
        color: var(--muted);
        font-family: var(--font-body);
        font-size: 15px;
        line-height: 1.8;
    ">
        Upload a PDF or DOCX above to activate the chat interface.<br>
        <span style="font-family: var(--font-mono); font-size: 12px; color: #3A4560;">
            Awaiting document...
        </span>
    </div>
    """, unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════════════════
#  FOOTER
# ════════════════════════════════════════════════════════════════════════════
st.markdown("""
<div class="footer-bar">
    Developed by Bibek Subedi &nbsp;|&nbsp; &copy; 2026 All Rights Reserved
</div>
""", unsafe_allow_html=True)